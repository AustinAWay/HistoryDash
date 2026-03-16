#!/usr/bin/env python3
"""Generate a coaching session log as a Word document."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Paths
OUTPUT_DIR = Path(r"C:\Users\Adam Work\PycharmProjects\coaching_coach\reports")
OUTPUT_DIR.mkdir(exist_ok=True)

# Colors
BLUE = RGBColor(37, 99, 235)
RED = RGBColor(220, 38, 38)
ORANGE = RGBColor(234, 88, 12)
GRAY = RGBColor(107, 114, 128)
WHITE = RGBColor(255, 255, 255)

# Student coaching schedules
STUDENTS = [
    {
        'name': 'Gus Castillo',
        'subject': 'AP Human Geography',
        'predicted': 1,
        'target': 3,
        'tier': 1,
        'sessions': [
            {'week': 1, 'focus': 'Diagnostic + Unit 5 kickoff'},
            {'week': 2, 'focus': 'Unit 5-6 completion check'},
            {'week': 3, 'focus': 'FRQ introduction + framework drilling'},
            {'week': 5, 'focus': 'FRQ intensive review'},
            {'week': 7, 'focus': 'Practice test debrief + final strategy'},
            {'week': 8, 'focus': 'Final confidence check'},
        ]
    },
    {
        'name': 'Branson Pfiester',
        'subject': 'AP Human Geography',
        'predicted': 2,
        'target': 4,
        'tier': 1,
        'accommodation': 'Double Time (+100%)',
        'sessions': [
            {'week': 1, 'focus': 'FRQ framework diagnostic'},
            {'week': 2, 'focus': 'Framework distinction practice'},
            {'week': 3, 'focus': 'Stimulus interpretation'},
            {'week': 5, 'focus': 'Timed FRQ with double-time'},
            {'week': 6, 'focus': 'Practice test review'},
            {'week': 8, 'focus': 'Final confidence check'},
        ]
    },
    {
        'name': 'Emma Cotner',
        'subject': 'AP World History',
        'predicted': 2,
        'target': 3,
        'tier': 1,
        'sessions': [
            {'week': 1, 'focus': 'Source analysis basics (HIPP)'},
            {'week': 2, 'focus': 'DBQ structure'},
            {'week': 4, 'focus': 'Units 8-9 review'},
            {'week': 6, 'focus': 'Practice test review'},
            {'week': 8, 'focus': 'Final prep'},
        ]
    },
    {
        'name': 'Saeed Tarawneh',
        'subject': 'AP World History',
        'predicted': 2,
        'target': 3,
        'tier': 1,
        'note': 'Also taking AP Psychology (May 12)',
        'sessions': [
            {'week': 1, 'focus': 'Priority unit mapping'},
            {'week': 2, 'focus': 'Content check + DBQ intro'},
            {'week': 4, 'focus': 'Practice test prep + source analysis'},
            {'week': 6, 'focus': 'Practice test review'},
            {'week': 7, 'focus': 'Final World History prep'},
        ]
    },
    {
        'name': 'Sydney Barba',
        'subject': 'AP Human Geography',
        'predicted': 3,
        'target': 4,
        'tier': 2,
        'note': 'FRQ at 17% - critical weakness',
        'sessions': [
            {'week': 1, 'focus': 'FRQ emergency intervention'},
            {'week': 2, 'focus': 'Framework identification check'},
            {'week': 4, 'focus': 'Timed FRQ review'},
            {'week': 7, 'focus': 'Final FRQ confidence check'},
        ]
    },
    {
        'name': 'Boris Dudarev',
        'subject': 'AP Human Geography',
        'predicted': 3,
        'target': 4,
        'tier': 2,
        'accommodation': 'Time +50%',
        'sessions': [
            {'week': 1, 'focus': 'Content prioritization'},
            {'week': 4, 'focus': 'Progress check + FRQ review'},
            {'week': 7, 'focus': 'Practice test debrief + accommodations prep'},
        ]
    },
    {
        'name': 'Zayen Szpitalak',
        'subject': 'AP Human Geography',
        'predicted': 3,
        'target': 5,
        'tier': 2,
        'sessions': [
            {'week': 1, 'focus': 'Spatial analysis + content plan'},
            {'week': 4, 'focus': 'FRQ practice'},
            {'week': 7, 'focus': 'Final prep'},
        ]
    },
    {
        'name': 'Aheli Shah',
        'subject': 'AP Human Geography',
        'predicted': 3,
        'target': 4,
        'tier': 2,
        'sessions': [
            {'week': 2, 'focus': 'Exam strategy'},
            {'week': 7, 'focus': 'Practice test review'},
        ]
    },
    {
        'name': 'Stella Cole',
        'subject': 'AP World History',
        'predicted': 3,
        'target': 4,
        'tier': 2,
        'sessions': [
            {'week': 1, 'focus': 'Content prioritization'},
            {'week': 4, 'focus': 'Progress check + DBQ introduction'},
            {'week': 7, 'focus': 'Practice test debrief + final strategy'},
        ]
    },
    {
        'name': 'Ella Dietz',
        'subject': 'AP World History',
        'predicted': 3,
        'target': 5,
        'tier': 2,
        'note': 'Best SAQ in group (7/9) - DBQ mechanics issue',
        'sessions': [
            {'week': 1, 'focus': 'DBQ formula'},
            {'week': 4, 'focus': 'DBQ practice review'},
            {'week': 7, 'focus': 'Final prep'},
        ]
    },
    {
        'name': 'Jackson Price',
        'subject': 'AP World History',
        'predicted': 3,
        'target': 4,
        'tier': 2,
        'note': 'Strongest MCQ (78%) - needs DBQ mechanics',
        'sessions': [
            {'week': 1, 'focus': 'DBQ structure'},
            {'week': 4, 'focus': 'DBQ practice review'},
            {'week': 7, 'focus': 'Final prep'},
        ]
    },
]

WEEK_DATES = {
    1: 'Mar 9-13',
    2: 'Mar 16-20',
    3: 'Mar 23-27',
    4: 'Mar 30 - Apr 3',
    5: 'Apr 6-10',
    6: 'Apr 13-17',
    7: 'Apr 27 - May 1',
    8: 'May 4-6',
}


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_row_shading(row, color_hex):
    """Set background color for entire row."""
    for cell in row.cells:
        set_cell_shading(cell, color_hex)


def add_student_section(doc, student):
    """Add a coaching log section for one student."""

    tier_color_hex = "DC2626" if student['tier'] == 1 else "EA580C"
    tier_color = RED if student['tier'] == 1 else ORANGE
    tier_label = "TIER 1: AT RISK" if student['tier'] == 1 else "TIER 2: BORDERLINE"

    # Student header
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    header_table.columns[0].width = Cm(6)
    header_table.columns[1].width = Cm(5)
    header_table.columns[2].width = Cm(5)

    row = header_table.rows[0]
    set_row_shading(row, tier_color_hex)

    # Name
    cell = row.cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(student['name'])
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(12)

    # Subject
    cell = row.cells[1]
    p = cell.paragraphs[0]
    run = p.add_run(student['subject'])
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

    # Tier
    cell = row.cells[2]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(tier_label)
    run.font.color.rgb = WHITE
    run.font.size = Pt(9)

    # Info line (score, accommodation, notes)
    info_parts = [f"Score {student['predicted']} → {student['target']}"]
    if student.get('accommodation'):
        info_parts.append(student['accommodation'])
    if student.get('note'):
        info_parts.append(student['note'])

    info_para = doc.add_paragraph()
    info_para.paragraph_format.space_before = Pt(0)
    info_para.paragraph_format.space_after = Pt(6)
    run = info_para.add_run(' | '.join(info_parts))
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # Sessions table
    num_sessions = len(student['sessions'])
    table = doc.add_table(rows=num_sessions + 1, cols=5)
    table.style = 'Table Grid'
    table.autofit = False

    # Column widths
    table.columns[0].width = Cm(2.5)   # Week
    table.columns[1].width = Cm(5)     # Focus
    table.columns[2].width = Cm(2.5)   # Coach
    table.columns[3].width = Cm(3)     # Date/Time
    table.columns[4].width = Cm(3)     # Notes

    # Header row
    header_row = table.rows[0]
    set_row_shading(header_row, "6B7280")
    headers = ['Week', 'Session Focus', 'Coach', 'Date/Time', 'Notes']
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)

    # Session rows
    for i, sess in enumerate(student['sessions'], 1):
        row = table.rows[i]
        week_num = sess['week']

        # Week cell
        cell = row.cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(f"Wk {week_num}")
        run.bold = True
        run.font.size = Pt(9)
        p.add_run(f"\n{WEEK_DATES[week_num]}").font.size = Pt(8)

        # Focus cell
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(sess['focus'])
        run.font.size = Pt(9)

        # Empty cells for coach, date/time, notes (leave blank for editing)
        for j in range(2, 5):
            cell = row.cells[j]
            # Set minimum height by adding empty content
            p = cell.paragraphs[0]
            p.add_run("")

        # Set row height
        row.height = Cm(1.2)

    # Add spacing after table
    doc.add_paragraph()


def build_doc():
    """Build the Word document."""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Title
    title = doc.add_heading('Coaching Session Log', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AP Coaching Program | Spring 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    # Week reference table
    doc.add_paragraph()
    ref_para = doc.add_paragraph()
    ref_para.add_run('Week Reference: ').bold = True

    ref_table = doc.add_table(rows=2, cols=8)
    ref_table.style = 'Table Grid'
    ref_table.autofit = False

    for i in range(8):
        ref_table.columns[i].width = Cm(2)

    # Header row
    for i in range(8):
        cell = ref_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(f"Wk {i+1}")
        run.bold = True
        run.font.size = Pt(8)
        set_cell_shading(cell, "2563EB")
        run.font.color.rgb = WHITE

    # Dates row
    for i in range(8):
        cell = ref_table.rows[1].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(WEEK_DATES[i+1])
        run.font.size = Pt(8)

    doc.add_paragraph()

    # Add each student
    for student in STUDENTS:
        add_student_section(doc, student)

    # Save
    doc_path = OUTPUT_DIR / "Coaching_Session_Log.docx"
    doc.save(str(doc_path))
    return doc_path


def main():
    print("Generating Coaching Session Log (Word)...")
    doc_path = build_doc()
    print(f"Done! Saved to: {doc_path}")


if __name__ == "__main__":
    main()
